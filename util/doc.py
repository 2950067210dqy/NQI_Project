import os
from docx import Document
from docx.enum.section import WD_SECTION_START
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
from docx.oxml import OxmlElement

# 指定文件夹路径
folder_path = '../'  # 替换为你的文件夹路径
output_docx = '动物箱软著所需代码v1.0.docx'  # 输出的 Word 文件名
excluded_folders = ['.venv', 'build','data','dist','log','model','other','resource','resource_py']  # 替换为你要排除的文件夹名称
# 创建一个新的 Word 文档
doc = Document()


def add_continuous_section(doc):
    """向文档中添加连续分节符"""
    """向文档中添加连续分节符"""
    # 仅更改下一节的类型为连续
    section = doc.add_section(start_type=WD_SECTION_START.CONTINUOUS)  # 代表连续分节符
    return section
# 遍历文件夹中的文件
for root, dirs, files in os.walk(folder_path):
    # 过滤掉要排除的文件夹
    dirs[:] = [d for d in dirs if d not in excluded_folders]
    # 使用 os.walk 递归遍历
    for filename in files:
        # 检查是否是.py文件，并且不是__init__.py文件
        if filename.endswith('.py') and filename not in ['__init__.py','doc.py'] :
            file_path = os.path.join(root, filename)  # 获取文件的完整路径
            print(file_path)  # 打印文件的完整路径
            try:
                with open(file_path, 'r', encoding='utf-8') as f:
                    # 读取文件内容并删除空行
                    code_lines = [line for line in f if line.strip()]
                    code_content = ''.join(code_lines)  # 将行拼接为单一字符串
            except Exception as e:
                print(f"无法读取文件: {file_path}，错误: {e}")
                continue

            # 将文件名添加为标题样式 (Heading 2) 并加粗
            title = doc.add_heading(level=2)
            run = title.add_run(f"{file_path.replace('../', '')}:")
            run.bold = True
            # 可选：设置标题对齐方式
            title.alignment = WD_PARAGRAPH_ALIGNMENT.LEFT



            # 添加代码内容到文档
            doc.add_paragraph(code_content)
            # 添加分节符
            add_continuous_section(doc)
# 保存文档
doc.save(output_docx)
print(f"文档已保存为 {output_docx}")